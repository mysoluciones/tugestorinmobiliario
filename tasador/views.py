import os
from django.shortcuts import render
from django.http import HttpResponse
from .forms import ContractForm, GastosForm, ReservaCompraForm
from django.conf import settings
from docx import Document
from io import BytesIO

def index(request):
    """Vista para la página de inicio."""
    return render(request, 'index.html')

def replace_placeholders(text, data):
    """Reemplaza los placeholders en el texto con los datos proporcionados."""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        text = text.replace(placeholder, str(value))
    return text

def generate_contract(request):
    """Vista para generar el contrato en formato DOCX."""
    if request.method == 'POST':
        form = ContractForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data

            # Cargar la plantilla DOCX
            template_path = os.path.join(settings.MEDIA_ROOT, 'contract_template.docx')
            doc = Document(template_path)

            # Reemplazar los placeholders en los párrafos
            for paragraph in doc.paragraphs:
                if any(f"{{{key}}}" in paragraph.text for key in data.keys()):
                    paragraph.text = replace_placeholders(paragraph.text, data)

            # Reemplazar los placeholders en los encabezados
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if any(f"{{{key}}}" in header.text for key in data.keys()):
                        header.text = replace_placeholders(header.text, data)

            # Reemplazar los placeholders en los pies de página
            for section in doc.sections:
                for footer in section.footer.paragraphs:
                    if any(f"{{{key}}}" in footer.text for key in data.keys()):
                        footer.text = replace_placeholders(footer.text, data)

            # Guardar el documento en un buffer
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Enviar el archivo al usuario
            response = HttpResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="contract.docx"'
            return response
    else:
        form = ContractForm()

    return render(request, 'generate_contract.html', {'form': form})

def generate_gastos_contract(request):
    """Vista para generar el contrato de gastos en formato DOCX."""
    if request.method == 'POST':
        form = GastosForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data

            # Cargar la plantilla DOCX
            template_path = os.path.join(settings.MEDIA_ROOT, 'gastos.docx')  # Asegúrate de que el nombre sea correcto
            doc = Document(template_path)

            # Reemplazar los placeholders en los párrafos
            for paragraph in doc.paragraphs:
                if any(f"{{{key}}}" in paragraph.text for key in data.keys()):
                    paragraph.text = replace_placeholders(paragraph.text, data)

            # Reemplazar los placeholders en los encabezados
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if any(f"{{{key}}}" in header.text for key in data.keys()):
                        header.text = replace_placeholders(header.text, data)

            # Reemplazar los placeholders en los pies de página
            for section in doc.sections:
                for footer in section.footer.paragraphs:
                    if any(f"{{{key}}}" in footer.text for key in data.keys()):
                        footer.text = replace_placeholders(footer.text, data)

            # Guardar el documento en un buffer
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Enviar el archivo al usuario
            response = HttpResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="gastos_contract.docx"'
            return response
    else:
        form = GastosForm()

    return render(request, 'generate_gastos_contract.html', {'form': form})


def generate_reserva_compra_contract(request):
    """Vista para generar el contrato de reserva de compra en formato DOCX."""
    if request.method == 'POST':
        form = ReservaCompraForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data

            # Cargar la plantilla DOCX
            template_path = os.path.join(settings.MEDIA_ROOT, 'reserva_compra.docx')  # Asegúrate de que el nombre sea correcto
            doc = Document(template_path)

            # Reemplazar los placeholders en los párrafos
            for paragraph in doc.paragraphs:
                if any(f"{{{key}}}" in paragraph.text for key in data.keys()):
                    paragraph.text = replace_placeholders(paragraph.text, data)

            # Reemplazar los placeholders en los encabezados
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if any(f"{{{key}}}" in header.text for key in data.keys()):
                        header.text = replace_placeholders(header.text, data)

            # Reemplazar los placeholders en los pies de página
            for section in doc.sections:
                for footer in section.footer.paragraphs:
                    if any(f"{{{key}}}" in footer.text for key in data.keys()):
                        footer.text = replace_placeholders(footer.text, data)

            # Guardar el documento en un buffer
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Enviar el archivo al usuario
            response = HttpResponse(doc_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="reserva_compra_contract.docx"'
            return response
    else:
        form = ReservaCompraForm()

    return render(request, 'generate_reserva_compra_contract.html', {'form': form})

def replace_placeholders(text, data):
    """Reemplaza los placeholders en el texto con los datos del formulario."""
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        text = text.replace(placeholder, str(value))
    return text
